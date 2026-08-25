import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    """Sets cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Applies clean subtle borders to a table."""
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def add_callout_box(doc, title, text, bg_color="F4F6F9", border_color="1B365D"):
    """Adds an executive callout box with a bold left accent border."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="28" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run_t = p.add_run(f"EXECUTIVE TAKEAWAY: {title.upper()}\n")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    run_b = p.add_run(text)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    run_b.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)

def build_document():
    doc = Document()
    
    # 1-inch margins all around
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("BNPL Economic, Regulatory & Litigation Assessment | Charles W. Kinslow IV, J.D., C.P.A.")
        f_run.font.name = "Calibri"
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    NAVY = RGBColor(0x1B, 0x36, 0x5D)
    SLATE = RGBColor(0x2B, 0x6C, 0xB0)
    CHARCOAL = RGBColor(0x22, 0x22, 0x22)
    MUTED = RGBColor(0x71, 0x80, 0x96)
    
    # Pre-header
    p_pre = doc.add_paragraph()
    p_pre.paragraph_format.space_before = Pt(0)
    p_pre.paragraph_format.space_after = Pt(2)
    r_pre = p_pre.add_run("EXECUTIVE FINANCIAL & LEGAL INTELLIGENCE REPORT")
    r_pre.font.name = "Calibri"
    r_pre.font.size = Pt(10)
    r_pre.bold = True
    r_pre.font.color.rgb = SLATE
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("Buy Now, Pay Later (BNPL) Sector Assessment:\nEconomic Fragility, 'Ghost Debt' Compounding, Regulatory Enforcement, and Active Litigation")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(17)
    r_title.bold = True
    r_title.font.color.rgb = NAVY
    
    # Byline
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(10)
    r_meta = p_meta.add_run("Authored by: Charles W. Kinslow IV, J.D., C.P.A.  |  Publication Date: August 2026  |  Coverage: Affirm (AFRM), Klarna, Afterpay, Zip, Sezzle")
    r_meta.font.name = "Calibri"
    r_meta.font.size = Pt(9.5)
    r_meta.italic = True
    r_meta.font.color.rgb = MUTED
    
    # Horizontal rule
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(0)
    p_hr.paragraph_format.space_after = Pt(10)
    p_hr_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="1B365D"/></w:pBdr>')
    p_hr._p.get_or_add_pPr().append(p_hr_border)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = SLATE
        return p

    def add_p(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.name = "Calibri"
            r_pre.font.size = Pt(10.5)
            r_pre.bold = True
            r_pre.font.color.rgb = CHARCOAL
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.italic = italic
        r.font.color.rgb = CHARCOAL
        return p

    def add_b(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.name = "Calibri"
            r_pre.font.size = Pt(10.5)
            r_pre.bold = True
            r_pre.font.color.rgb = CHARCOAL
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.color.rgb = CHARCOAL
        return p

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    add_h1("1. Executive Summary: The Fintech Illusion Meets Balance Sheet Reality")
    add_p(
        "Let’s cut through the fintech marketing deck. For the past six years, the Buy Now, Pay Later (BNPL) sector sold Wall Street and consumers on a convenient fairy tale: that wrapping a four-installment microloan inside a slick mobile UI somehow represented a virtuous, high-tech democratization of credit. They claimed it was completely distinct from the greedy, predatory revolving credit card model. No late fees, no compounding interest, no friction. Just seamless point-of-sale budgeting for Gen Z and millennials."
    )
    add_p(
        "Fast forward to 2026, and the mathematics of reality have caught up. The zero-interest-rate environment that funded this hyper-growth circus is long gone. When wholesale cost of capital rose, the pure-play, zero-fee 'Pay-in-4' model immediately exposed its fatal structural defect: you cannot build a sustainable, multibillion-dollar financial institution on sub-2% merchant discount fees while writing uncollateralized loans to consumers with zero underwriting and zero income verification."
    )
    add_p(
        "What we are witnessing across the industry today is a sharp bifurcation. The operators who survived—chiefly Affirm—did so not by staying pure to the Pay-in-4 ethos, but by quietly morphing into high-APR installment lenders (charging up to 36% interest) and securing captive distribution through Apple Pay and Amazon. Meanwhile, pure-play pioneers have either collapsed into liquidation (Laybuy, Openpay), retreated into geographical bunkers (Zip Co dumping its global footprint), or suffered massive accounting bloodbaths (Klarna posting a -$273M loss in 2025 as credit loss provisioning caught up to reality)."
    )
    add_p(
        "Simultaneously, the regulatory arbitrage that allowed these companies to operate in the shadows is being dismantled on every front. From New York's aggressive BNPL Act rulemaking to the CFPB's Regulation Z card-issuer classification and mounting federal class actions targeting predatory overdraft fee causation, the era of frictionless, unregulated lending is definitively over."
    )

    add_callout_box(
        doc,
        "The Bottom Line on BNPL Economics",
        "BNPL companies aren't charities, and the laws of credit underwriting are immutable. When you hand out unmonitored installment loans without checking ability to repay, you don't eliminate credit risk—you merely hide it off the credit bureau radar. Now, with 47% of users missing payments and 29% financing everyday groceries, the compounding weight of 'ghost debt' is squeezing both consumer balance sheets and fintech margins."
    )

    # --- SECTION 2: MACROECONOMIC CRACKDOWN & GHOST DEBT ---
    add_h1("2. Macroeconomic Fragility: 'Ghost Debt', Multi-App Stacking, and the Grocery Mirage")
    add_p(
        "If you want to understand the true health of the American consumer, stop looking at polished Bureau of Economic Analysis top-line spending figures and look at how people are actually settling their grocery carts. The empirical reality of the BNPL market in 2025–2026 tells a story of acute subprime distress masquerading as fintech adoption:"
    )
    add_b("47% Delinquency Rate: A staggering 47% of active BNPL users reported missing at least one payment over the trailing twelve months—up from 41% in 2025 and 34% in 2024. The Silicon Valley myth that short installment terms naturally enforce borrower discipline has completely unraveled.", "Exploding Consumer Delinquencies: ")
    add_b("63% Multi-App Loan Stacking: Roughly two-thirds of active users hold concurrent installment loans across Affirm, Klarna, Afterpay, and Zip simultaneously. Over 25% carry three or more active plans at once. Because each app underwrites in total isolation, borrowers easily stack repayment schedules that consume 100% of their weekly liquidity.", "Rampant Debt Stacking: ")
    add_b("29% of Loans Are for Groceries & Survival: In 2024, groceries and household staples accounted for 14% of BNPL volume. By mid-2026, that number surged to nearly 30%. More than 54% of borrowers openly admit they use BNPL products because they simply lack the liquid cash to make ends meet between paychecks. This isn't discretionary lifestyle financing; it is high-tech payday lending for basic nutrition.", "Financing the Dinner Table: ")
    add_b("The 'Ghost Debt' Blind Spot: Because BNPL originations rely almost exclusively on soft credit checks and systematically bypass traditional credit reporting repositories, billions in short-term installment debt remain completely invisible to mortgage lenders, auto finance desks, and credit card underwriters. The Richmond Federal Reserve and Fitch Ratings have repeatedly warned that this shadow leverage distorts debt-to-income (DTI) metrics and masks systemic credit fragility.", "Credit Bureau Invisibility: ")

    # Table 1: Comparative Performance
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(8)
    p_t.paragraph_format.space_after = Pt(3)
    r_tl = p_t.add_run("Table 1: BNPL Market Participants — Reported Metrics vs. Balance Sheet Reality (2024–2026)")
    r_tl.font.name = "Calibri"
    r_tl.font.size = Pt(10)
    r_tl.bold = True
    r_tl.font.color.rgb = NAVY

    tbl_perf = doc.add_table(rows=6, cols=4)
    tbl_perf.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_perf.autofit = False
    set_table_borders(tbl_perf)

    headers = ["Participant", "Reported Financial Health", "Credit Quality & Loss Provisioning", "Strategic Reality & Legal Exposure"]
    col_widths = [Inches(1.4), Inches(1.6), Inches(1.7), Inches(1.8)]
    
    for i, title in enumerate(headers):
        cell = tbl_perf.cell(0, i)
        cell.width = col_widths[i]
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    table_rows = [
        ("Affirm Holdings\n(NASDAQ: AFRM)", "~$36B FY25 GMV;\nApproaching GAAP operating profitability", "30+ day delinquencies at 2.7%; Pay-in-4 loss <1%; charge-offs ~3.5%", "Thrives by moving away from Pay-in-4 into 36% APR loans; relies on Apple/Amazon checkout moats."),
        ("Klarna Bank AB", "$3.5B revenue (+25% YoY),\nbut -$273M net loss in 2025; +$1M in Q1 2026", "Crushed by upfront CECL/IFRS 9 credit loss provisions on long-term loan cohorts", "Valuation hammered from $45.6B peak; facing N.D. Ill. predatory lending class actions and EU suits."),
        ("Afterpay\n(Block, Inc.)", ">$1.04B annual revenue;\nIntegrated into Cash App", "Maintains optics by forcing automated debit pulls on depleted checking accounts", "Defending class actions (Sands v. Afterpay) for triggering devastating bank overdraft/NSF fee cascades."),
        ("Zip Co Limited\n(ASX: ZIP)", "$79.9M FY25 NPAT\n(+1,110% YoY turnaround via radical restructuring)", "Enforces tight automated limits; reliant on U.S. consumer cash flow (>80% earnings)", "Fled UK, European, and global markets to avoid mounting regulatory compliance costs."),
        ("Laybuy & Openpay\n(The Graveyard)", "Laybuy: Receivership June 2024\nOpenpay: Insolvent liquidation", "Total credit collapse; unable to refinance wholesale warehouse debt facilities", "Conclusive mathematical proof that standalone, zero-fee Pay-in-4 cannot survive at scale.")
    ]

    for row_idx, row_content in enumerate(table_rows, start=1):
        bg = "F4F6F9" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_content):
            cell = tbl_perf.cell(row_idx, col_idx)
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(8.5)
            r.font.color.rgb = CHARCOAL

    p_post = doc.add_paragraph()
    p_post.paragraph_format.space_before = Pt(4)
    p_post.paragraph_format.space_after = Pt(6)

    # --- SECTION 3: PROVIDER DEEP DIVE ---
    add_h1("3. Individual Provider Breakdown: From Product Ingenuity to Regulatory Defense")

    add_h2("A. Affirm Holdings, Inc. (AFRM): The Hybrid Installment Machine")
    add_p(
        "Affirm is often cited as the gold standard of the sector, and for good reason: CEO Max Levchin recognized early on that pure Pay-in-4 was an economic dead end. While peers doubled down on fee-free microloans, Affirm built an underwriting engine geared around longer-term (up to 36-month) interest-bearing installment loans carrying APRs up to 36%."
    )
    add_b("Credit Metrics: Affirm maintains an allowance for credit losses of ~5.4% on loans held for investment. Its 30-plus-day delinquencies have held stable near 2.7%, with net charge-offs at ~3.5% on installment cohorts and Pay-in-4 losses contained below 1% of GMV.", "Underwriting Discipline: ")
    add_b("The Funding Reality: Affirm relies heavily on warehouse facilities and Asset-Backed Securitization (ABS) debt. While narrowing ABS spreads and lower benchmark rates have supported its Revenue Less Transaction Costs (RLTC) margins, Affirm remains highly exposed to capital market shocks.", "Capital Markets Exposure: ")
    add_b("The Apple & Amazon Distribution Moat: When Apple shuttered Apple Pay Later in June 2024, Affirm stepped into the vacuum, embedding itself directly into Apple Pay and Amazon. This gives Affirm unparalleled scale, but makes it captive to high merchant acquisition costs and merchant contract renegotiations.", "Distribution Leverage: ")

    add_h2("B. Klarna Bank AB: The Costly Transition to Digital Banking")
    add_p(
        "Klarna’s financial reports read like a cautionary tale of fintech growing pains. After celebrating a modest $21 million net profit in 2024, Klarna plunged into a -$273 million net loss for FY2025. The culprit wasn't top-line stagnation—revenue actually rose 25% to $3.5 billion. The culprit was credit accounting reality: as Klarna pushed aggressively into multi-month 'Fair Financing' loans and banking cards, CECL and IFRS 9 accounting rules forced it to record massive upfront provisions for expected credit losses."
    )
    add_p(
        "Klarna scraped together a $1 million net profit in Q1 2026 only by ruthlessly slashing operational headcount through AI automation and hiking merchant take-rates. But its valuation remains a shadow of its $45.6 billion pandemic peak, and its underwriting models are facing severe scrutiny in both US and European courts."
    )

    add_h2("C. Market Failures: The Graveyard of Standalone Pay-in-4")
    add_p(
        "The economic carnage among weaker entrants demonstrates the unviability of operating without immense scale or diversified lending products:"
    )
    add_b("Laybuy Group: Placed into receivership on June 17, 2024, after collapsing under high consumer default rates and failing to secure emergency debt funding. Klarna acquired the customer base, while liquidators continue untangling creditor losses in New Zealand High Court.", "Laybuy Receivership & Liquidation: ")
    add_b("Openpay: Insolvent and liquidated after burning through capital on failed US/UK expansion plans, leaving merchants and warehouse lenders holding empty bags.", "Openpay Collapse: ")
    add_b("Apple Pay Later: In June 2024, Apple—with hundreds of billions in cash—officially discontinued its first-party lending unit (Apple Financing LLC). Apple decided that managing credit loss provisions, collections, and state lending compliance was a value-destroying headache, choosing instead to outsource the risk to Affirm.", "Apple's Surrender: ")

    # --- SECTION 4: THE REGULATORY REGIME ---
    add_h1("4. The Tightening Regulatory Web: Federal, State, and Global Governance")
    add_p(
        "For years, BNPL lenders exploited a glaring statutory loophole: the Truth in Lending Act (TILA) and Regulation Z traditionally defined a consumer credit transaction based on whether it involved 'more than four installments' or carried a 'finance charge' (15 U.S.C. § 1602(g)). By structuring loans as zero-interest 4-pay transactions, BNPL firms claimed they were completely exempt from credit card rules. That loophole is now closing on every front."
    )

    add_h2("A. Federal Regulatory Framework & Proposed Legislation")
    add_bullet = add_b
    add_bullet("CFPB Regulation Z Interpretive Rule (May 2024): The CFPB issued a formal rule designating BNPL lenders as 'card issuers' and 'creditors' under Regulation Z (12 C.F.R. Part 1026). This subjects BNPL apps to: (1) mandatory 60-day billing error investigation and resolution rights; (2) mandatory merchant refund crediting obligations; and (3) periodic billing statement delivery.", "CFPB Card-Issuer Classification: ")
    add_bullet("Buy Now, Pay Later Consumer Protection Act (June 2026): Introduced by Rep. Dan Goldman (D-NY) to formally establish statutory BNPL protections at the federal level, removing any ambiguity regarding APR disclosures, fee stacking, and consumer dispute rights.", "Federal Statutory Action: ")

    add_h2("B. State Enforcement: New York & California")
    add_bullet("New York BNPL Act (NYDFS Proposed Rulemaking — 3 NYCRR Part 425): Enacted in May 2025; NYDFS issued proposed rules with comments closing September 14, 2026. The regulations impose: (1) mandatory state licensing; (2) application of New York’s 16% civil usury interest cap; (3) strict limits on late fees and stacking; (4) standardized dispute resolution; and (5) express consumer opt-in consent before fintechs can monetize user transaction data.", "New York State Drops the Hammer: ")
    add_bullet("California DFPI & AB 2350: California regulates BNPL products as loans under the California Financing Law (CFL) and aggressively pursues misleading marketing under the CCFPL, while moving AB 2350 to restrict predatory 'Rent Now, Pay Later' schemes.", "California Closes Loopholes: ")

    add_h2("C. International Regulatory Crackdowns")
    add_bullet("European Union Consumer Credit Directive II (Directive (EU) 2023/2225): Completely eliminated the under-€200 exemption, forcing all BNPL lenders to perform robust creditworthiness checks and provide standardized pre-contractual consumer disclosures.", "EU CCD II: ")
    add_bullet("United Kingdom Financial Conduct Authority (FCA): HM Treasury and the FCA are bringing BNPL under the Financial Services and Markets Act (FSMA), subjecting providers to Section 21 financial promotion rules, affordability assessments, and Financial Ombudsman Service (FOS) oversight.", "UK FCA Regime: ")
    add_bullet("Australia National Consumer Credit Protection (NCCP) Act: Requires all BNPL operators to hold an Australian Credit Licence (ACL) and comply with statutory responsible lending and unsuitability mandates.", "Australia NCCP Reforms: ")

    # --- SECTION 5: LITIGATION FRONT ---
    add_h1("5. Active Litigation Front: Trade Association Challenges, Overdraft Class Actions, and Breaches")
    add_p(
        "As regulatory walls close in, the battleground has shifted into federal and state courts. The litigation falls into three distinct categories:"
    )

    add_h2("A. Industry vs. Regulators: Financial Technology Association v. CFPB")
    add_p(
        "In October 2024, the Financial Technology Association (FTA)—representing Affirm, Klarna, Block/Afterpay, and Zip—filed a landmark lawsuit against the CFPB in the U.S. District Court for the District of Columbia (Case 1:24-cv-02945 / 1:24-cv-02966).",
        "The Lawsuit: "
    )
    add_b("APA Notice-and-Comment Claim: The FTA argues the CFPB violated the Administrative Procedure Act (5 U.S.C. § 553) by issuing a binding legislative rule masquerading as an 'interpretive rule' without undergoing mandatory public notice-and-comment.", "Administrative Procedure Act Violation: ")
    add_b("Statutory Overreach under TILA: Plaintiffs assert the CFPB exceeded its statutory authority under 15 U.S.C. § 1601 et seq. by contorting the definition of 'credit card' to encompass closed-end digital installment accounts.", "Exceeding TILA Authority: ")
    add_b("Operational Incompatibility: The complaint claims that credit card billing dispute cycles and monthly statement rules are structurally incompatible with 6-week microloan lifecycles.", "Operational Mismatch: ")

    add_h2("B. Consumer Class Actions: Overdrafts and Inability-to-Pay")
    add_b("Klarna Predatory Lending Class Action (N.D. Ill. 2026): A putative class action asserting that Klarna extends unvetted revolving credit ('Purchasing Power') without verifying income or ability to repay, relying on forced automatic debit withdrawals that systematically trigger cascading $35 bank overdraft and NSF fees.", "Klarna (N.D. Ill. 2026): ")
    add_b("Afterpay Overdraft Deception Litigation (Sands v. Afterpay US, Inc.): Class actions exposing the predatory reality behind 'no fees, no interest' marketing, demonstrating that uncoordinated recurring ACH debits systematically drain consumer checking accounts.", "Afterpay Overdraft Suits: ")
    add_b("Klarna Netherlands Collective Action (July 2026): A Dutch consumer protection foundation taking Klarna to court for systemic violations of EU credit assessment laws and unlawful penalty fee structures.", "European Collective Action: ")

    add_h2("C. Data Security & Arbitration Realities")
    add_b("Affirm & Evolve Bank Data Breach Settlement (MDL 2:24-md-03127): A class settlement finalized in December 2025 after millions of Affirm borrower records were compromised via partner bank Evolve Bank & Trust.", "Affirm / Evolve Bank Settlement: ")
    add_b("Affirm Mandatory Arbitration Enforcement: Affirm has successfully deflected class litigation regarding effective merchant interest markups and return disputes by enforcing binding individual arbitration clauses in its Terms of Service.", "Affirm Consumer Arbitration: ")

    # --- SECTION 6: TABLE OF AUTHORITIES ---
    add_h1("6. Comprehensive Table of Authorities & Legal Citations")

    add_h2("A. Statutes, Regulations & Legislative Materials")
    add_b("Consumer Financial Protection Bureau, Interpretive Rule: Use of Digital User Accounts to Access Buy Now, Pay Later Products, 89 Fed. Reg. 47072 (May 31, 2024) (codified at 12 C.F.R. Part 1026 / Regulation Z).", "1. ")
    add_b("Truth in Lending Act (TILA), 15 U.S.C. § 1601 et seq.; 15 U.S.C. § 1602(g) (creditor definition); Regulation Z, 12 C.F.R. § 1026.2(a)(15) (card issuer definitions).", "2. ")
    add_b("Administrative Procedure Act (APA), 5 U.S.C. §§ 553 (rulemaking), 701(a) (agency discretion), 706(2)(A) (arbitrary and capricious standard).", "3. ")
    add_b("Buy Now, Pay Later Consumer Protection Act, H.R. 8762, 119th Cong. (introduced by Rep. Daniel S. Goldman, June 2026).", "4. ")
    add_b("New York Buy Now, Pay Later Act, N.Y. Banking Law Art. 9-B §§ 380–395 (enacted May 2025); New York Department of Financial Services (NYDFS), Proposed 3 NYCRR Part 425 (Notice of Proposed Rulemaking, July 2026; public comments due September 14, 2026).", "5. ")
    add_b("California Financing Law (CFL), Cal. Fin. Code § 22000 et seq.; California Consumer Financial Protection Law (CCFPL), Cal. Fin. Code § 90000 et seq.; Cal. Assemb. Bill 2350 (2025–2026 Reg. Sess.).", "6. ")
    add_b("Directive (EU) 2023/2225 of the European Parliament and of the Council of 18 October 2023 on credit agreements for consumers (CCD II).", "7. ")
    add_b("United Kingdom Financial Services and Markets Act 2000 (FSMA), c. 8; HM Treasury & FCA BNPL Consultation & Draft Regulations (2024–2026).", "8. ")
    add_b("Treasury Laws Amendment (Buy Now Pay Later) Bill 2024 (amending National Consumer Credit Protection Act 2009 (Cth), Australia).", "9. ")

    add_h2("B. Case Law, Dockets & Judicial Proceedings")
    add_b("Financial Technology Association v. Consumer Financial Protection Bureau, No. 1:24-cv-02945 (D.D.C. filed Oct. 18, 2024) (APA challenge to Regulation Z Interpretive Rule).", "10. ")
    add_b("Doe et al. v. Klarna Inc., No. 1:26-cv-02184 (N.D. Ill. filed 2026) (class action asserting predatory lending, lack of underwriting, and automated overdraft causation).", "11. ")
    add_b("Sands v. Afterpay US, Inc., No. 4:21-cv-02758 (N.D. Cal.) (class action asserting deceptive marketing and unlawful ACH overdraft fee triggers).", "12. ")
    add_b("In re Evolve Bank & Trust Customer Data Security Breach Litigation, No. 2:24-md-03127 (W.D. Tenn. final settlement approved Dec. 2025).", "13. ")
    add_b("Stichting Consumentenbelang v. Klarna Bank AB (District Court of Amsterdam, filed July 2026) (Dutch collective action on creditworthiness assessments).", "14. ")

    add_h2("C. Financial Reports, SEC Filings & Macroeconomic Studies")
    add_b("Affirm Holdings, Inc., Form 10-K for FY ended June 30, 2025 & Form 10-Q for Period Ended March 31, 2026 (disclosing ~$36B GMV, 5.4% credit loss allowance, 2.7% 30-day delinquency).", "15. ")
    add_b("Klarna Bank AB (publ), Annual Report 2024 & Interim Report January–December 2025 ($21M net income 2024; -$273M net loss 2025; +$1M net profit Q1 2026).", "16. ")
    add_b("Block, Inc., Form 10-K for Year Ended December 31, 2025 (disclosing Afterpay revenue contributions >$1.04B and Cash App integration).", "17. ")
    add_b("Zip Co Limited, FY25 Annual Financial Report (ASX: ZIP) ($79.9M NPAT; divestment of UK/European operations).", "18. ")
    add_b("Federal Reserve Bank of Richmond, 'Buy Now, Pay Later: A Snapshot of Market Trends and Consumer Risks,' Economic Brief (2025/2026).", "19. ")
    add_b("Fitch Ratings, 'Global Consumer Credit Outlook: Shadow Debt and BNPL Delinquency Pressures' (Late 2025/2026 Report).", "20. ")
    add_b("Laybuy Group Holdings Ltd (in Receivership & Liquidation), Liquidators' Six-Monthly Report (High Court of New Zealand, 2024–2025).", "21. ")

    # Closing Signature Block
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(16)
    p_sig.paragraph_format.space_after = Pt(2)
    r_sig = p_sig.add_run("— Charles W. Kinslow IV, J.D., C.P.A.")
    r_sig.font.name = "Calibri"
    r_sig.font.size = Pt(11)
    r_sig.bold = True
    r_sig.font.color.rgb = NAVY
    
    p_tags = doc.add_paragraph()
    p_tags.paragraph_format.space_before = Pt(0)
    p_tags.paragraph_format.space_after = Pt(6)
    r_tags = p_tags.add_run("@Affirm @CFPB @Financial Technology Association @Morgan Lewis & Bockius LLP #BNPL #BuyNowPayLater #ConsumerProtection #Fintech #GhostDebt #RegulatoryArchive2026")
    r_tags.font.name = "Calibri"
    r_tags.font.size = Pt(9.5)
    r_tags.italic = True
    r_tags.font.color.rgb = MUTED

    target_filename = "BNPL_Industry_Economic_Regulatory_and_Litigation_Report_2026.docx"
    output_path = os.path.join(r"c:\Users\Charwiz43\.gemini\antigravity\scratch\Affirm\regulatory-archive-2026", target_filename)
    
    try:
        doc.save(output_path)
        print(f"SUCCESS: Overwrote target file at: {output_path}")
    except PermissionError:
        print(f"FILE_LOCKED: {output_path} is currently locked by Word.")
        fallback_path = os.path.join(r"c:\Users\Charwiz43\.gemini\antigravity\scratch\Affirm\regulatory-archive-2026", "BNPL_Industry_Economic_Regulatory_and_Litigation_Report_2026_Updated.docx")
        doc.save(fallback_path)
        print(f"SAVED_FALLBACK: {fallback_path}")

if __name__ == "__main__":
    build_document()
